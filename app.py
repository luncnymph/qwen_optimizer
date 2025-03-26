from flask import Flask, request, send_file, render_template_string, jsonify
import requests
import os
from docx import Document
from werkzeug.utils import secure_filename
import shutil
import time
import langdetect
import logging
import uuid

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)

HTML_TEMPLATE = '''
<!doctype html>
<html lang="en">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1, shrink-to-fit=no">
    <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.3.1/css/bootstrap.min.css">
    <title>Qwen Document Optimizer</title>
    <style>
        body { background-color: #343a40; font-family: Arial, sans-serif; color: #f8f9fa; }
        .container { max-width: 700px; margin-top: 60px; background: #495057; padding: 40px; border-radius: 12px; }
        .logo { display: block; margin: 0 auto 30px; max-width: 180px; }
        h2 { color: #17a2b8; text-align: center; font-weight: bold; margin-bottom: 30px; }
        .form-group label { font-weight: 600; color: #f8f9fa; }
        .custom-file-upload { 
            display: inline-block; padding: 8px 20px; cursor: pointer; 
            background-color: #17a2b8; color: white; border-radius: 6px;
        }
        #file { display: none; }
        #file-name { margin-left: 15px; color: #adb5bd; font-style: italic; }
        .btn-primary { width: 100%; padding: 12px; background-color: #17a2b8; border: none; }
        .api-key-guide { margin-top: 20px; font-size: 0.9em; color: #adb5bd; }
    </style>
</head>
<body>
<div class="container">
    <img src="{{ url_for('static', filename='logo.png') }}" alt="Logo" class="logo">
    <h2>Optimize Document with Qwen AI</h2>
    <form id="upload-form" enctype="multipart/form-data">
        <div class="form-group">
            <label for="api_key">Qwen API Key:</label>
            <input type="text" class="form-control" id="api_key" name="api_key" required>
            <input type="hidden" id="api_key_hidden" name="api_key_hidden">
        </div>       
        <div class="form-group">
            <label for="file">Select Word Document (.docx):</label><br>
            <label class="custom-file-upload">
                Choose File
                <input type="file" id="file" name="file" accept=".docx" required onchange="updateFileName()">
            </label>
            <span id="file-name">No file selected</span>
        </div>
        <button type="submit" class="btn btn-primary" id="submit-btn">Submit</button>
    </form>
    <div id="message" class="alert alert-info mt-3" style="display: none;"></div>
    <div class="api-key-guide">
        <h4>How to obtain Qwen API key</h4>
        <ol>
            <li><strong>Sign Up or Log In to Alibaba Cloud</strong><br>
                Visit the <a href="https://www.alibabacloud.com/" target="_blank">Alibaba Cloud website</a> and create an account if you don't already have one. If you already have an account, log in using your credentials.</li>
            <li><strong>Access the Qwen Dashboard</strong><br>
                Once logged in, navigate to the <a href="https://dashscope.console.aliyun.com/" target="_blank">Qwen dashboard</a>. This is where you can manage your Qwen models and access related services.</li>
            <li><strong>Create an API Key</strong><br>
                In the Qwen dashboard, go to the API Keys section. Click on the option to <em>Create API Key</em> or <em>Generate API Key</em>. Follow the prompts to generate a new API key. Make sure to copy and save the API key in a secure place, as it will not be displayed again after you leave the page.</li>
        </ol>
    </div>
</div>
<script src="https://code.jquery.com/jquery-3.3.1.min.js"></script>
<script>
    function updateFileName() {
        const fileInput = document.getElementById('file');
        const fileNameDisplay = document.getElementById('file-name');
        fileNameDisplay.textContent = fileInput.files[0]?.name || 'No file selected';
    }

    function maskApiKey(value) {
        if (value.length <= 8) return value; // Show short keys as-is
        const firstPart = value.substring(0, 4);
        const lastPart = value.substring(value.length - 4);
        const maskedPart = '*'.repeat(value.length - 8);
        return firstPart + maskedPart + lastPart;
    }

    $(document).ready(function() {
        const $apiKeyInput = $('#api_key');
        const $apiKeyHidden = $('#api_key_hidden');

        // Mask the API key on input
        $apiKeyInput.on('input', function() {
            const realValue = this.value;
            $apiKeyHidden.val(realValue); // Store real value in hidden input
            this.value = maskApiKey(realValue); // Display masked value
        });

        // Preserve real value on focus out
        $apiKeyInput.on('focus', function() {
            const realValue = $apiKeyHidden.val();
            if (realValue) this.value = realValue; // Show real value when focused
        });

        $apiKeyInput.on('blur', function() {
            const realValue = $apiKeyHidden.val();
            if (realValue) this.value = maskApiKey(realValue); // Mask again when unfocused
        });

        $('#upload-form').on('submit', function(e) {
            e.preventDefault();
            const $submitBtn = $('#submit-btn');
            const $message = $('#message');
            $submitBtn.text('Processing...Please wait').prop('disabled', true);
            $message.hide();

            const formData = new FormData(this);
            formData.set('api_key', $apiKeyHidden.val() || $apiKeyInput.val()); // Use real value

            $.ajax({
                url: '/upload',
                type: 'POST',
                data: formData,
                processData: false,
                contentType: false,
                xhr: function() {
                    const xhr = new window.XMLHttpRequest();
                    xhr.responseType = 'blob';
                    return xhr;
                },
                success: function(data, status, xhr) {
                    const filename = xhr.getResponseHeader('Content-Disposition')?.match(/filename="(.+)"/)?.[1] || 'optimized_document.docx';
                    const blob = new Blob([data], { type: 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' });
                    const url = window.URL.createObjectURL(blob);
                    const link = document.createElement('a');
                    link.href = url;
                    link.download = filename;
                    document.body.appendChild(link);
                    link.click();
                    document.body.removeChild(link);
                    window.URL.revokeObjectURL(url);

                    $submitBtn.text('Submit').prop('disabled', false);
                },
                error: function(xhr) {
                    let errorMsg = 'An error occurred';
                    if (xhr.responseJSON && xhr.responseJSON.message) {
                        errorMsg = xhr.responseJSON.message;
                    }
                    $message.text(errorMsg).show();
                    $submitBtn.text('Submit').prop('disabled', false);
                }
            });
        });
    });
</script>
</body>
</html>
'''

UPLOAD_FOLDER = '/tmp/uploads'
OUTPUT_FOLDER = '/tmp/outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

QWEN_API_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'docx'

def detect_language(text):
    try:
        return langdetect.detect(text)
    except:
        return 'en'

def make_api_request(headers, data, max_retries=3):
    for attempt in range(max_retries):
        try:
            response = requests.post(
                QWEN_API_URL,
                headers=headers,
                json=data,
                timeout=120,
                proxies={"http": None, "https": None}
            )
            logger.debug(f"API Response Status: {response.status_code}")
            logger.debug(f"API Response Content: {response.text}")
            response.raise_for_status()
            return response
        except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as e:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                time.sleep(wait_time)
                continue
            raise requests.exceptions.ReadTimeout(f"API request timed out after {max_retries} attempts: {str(e)}")
        except requests.exceptions.RequestException as e:
            raise e

@app.route('/', methods=['GET'])
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/upload', methods=['POST'])
def upload_file():
    if not all(k in request.form for k in ('api_key',)) or 'file' not in request.files:
        return jsonify({'message': 'Please provide API key and document'}), 400

    file = request.files['file']
    api_key = request.form['api_key']

    if not allowed_file(file.filename):
        return jsonify({'message': 'Only .docx files are supported'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    try:
        doc = Document(filepath)
        text_content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

        if not text_content.strip():
            return jsonify({'message': 'Document is empty'}), 400

        lang = detect_language(text_content)
        is_chinese = lang.startswith('zh')

        if is_chinese:
            prompt = (
                "我有一个中文Word文档，希望优化其清晰度、可读性和专业性。请审查文本并在必要时进行改进，确保语言简洁明了，提升文档的整体流畅性。输出中不要包含#或*等符号，提供干净、专业格式的优化文本：\n\n"
                f"{text_content}"
            )
        else:
            prompt = (
                "I have a Word document that I would like to optimize for clarity, readability, and professionalism. "
                "Please review the text and make improvements where necessary. Ensure that the language is clear and concise, "
                "and enhance the overall flow of the document. Do not include any symbols such as # or * in the output. "
                "Provide the optimized text in a clean, professional format:\n\n"
                f"{text_content}"
            )

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json; charset=utf-8"
        }
        data = {
            "model": "qwen-long",
            "messages": [{"role": "user", "content": prompt}],
            "parameters": {"max_tokens": 8192, "temperature": 0.7}
        }

        response = make_api_request(headers, data)

        result = response.json()
        optimized_text = result.get('choices', [{}])[0].get('message', {}).get('content', '').strip()
        optimized_text = optimized_text.replace('*', '').replace('#', '')

        if not optimized_text:
            return jsonify({'message': 'No content received from API'}), 500

        new_doc = Document()
        for paragraph in optimized_text.split('\n'):
            if paragraph.strip():
                new_doc.add_paragraph(paragraph.strip())
        output_path = os.path.join(OUTPUT_FOLDER, f'optimized_{filename}')
        new_doc.save(output_path)

        if not os.path.exists(output_path):
            raise Exception("Output file was not created")

        response = send_file(
            output_path,
            as_attachment=True,
            download_name=f'optimized_{filename}',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        if os.path.exists(filepath):
            os.remove(filepath)
        if os.path.exists(output_path):
            os.remove(output_path)
        return response

    except requests.exceptions.ProxyError:
        return jsonify({'message': 'Proxy connection failed - check network settings'}), 500
    except requests.exceptions.ReadTimeout:
        return jsonify({'message': 'Request timed out after multiple attempts - try a smaller document or check network stability'}), 500
    except requests.exceptions.HTTPError as e:
        status_code = e.response.status_code
        logger.error(f"HTTP Error {status_code}: {e.response.text}")
        message = {
            400: 'Invalid request parameters',
            401: 'Invalid API key',
            403: 'Access forbidden - check API key permissions or contact Qwen support',
            404: 'API endpoint not found',
            500: 'Server error - try again later'
        }.get(status_code, f'API error {status_code}')
        return jsonify({'message': message}), status_code
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        return jsonify({'message': f'System error: {str(e)}'}), 500

if __name__ == '__main__':
    try:
        app.run(debug=True)
    finally:
        shutil.rmtree(UPLOAD_FOLDER, ignore_errors=True)
        shutil.rmtree(OUTPUT_FOLDER, ignore_errors=True)